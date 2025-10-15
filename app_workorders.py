# app_workorders.py — Parquet fast-path drop‑in + CrossRef + Service Procedures (LOCAL)
# --------------------------------------------------------------
# SPF Work Orders (reads local Workorders.xlsx by default; Parquet cache)
# Pages: Asset History • Work Orders • Cross Reference • Service Report • Service History • Service Procedure Filter
# Privacy-safe by Location; Dates normalized; “Data last updated” (ET)
# Requires: pandas, pyarrow, xlsxwriter, streamlit-authenticator, openpyxl
# (Optional) python-docx for Word exports from app buttons
# --------------------------------------------------------------

from __future__ import annotations
import io, re, os, time, platform, tempfile
from urllib.parse import quote
from pathlib import Path
from collections.abc import Mapping
from datetime import datetime, timedelta, timezone

import pandas as pd
import streamlit as st
import yaml

APP_VERSION = "2025.10.15p5-local"

# ---------- small CSS: hide Streamlit chrome / shrink controls ----------
st.set_page_config(page_title="SPF Work Orders", page_icon="🧰", layout="wide")
st.markdown(
    """
    <style>
      /* Hide viewer badge / deploy button / footer */
      .stDeployButton, footer, header {visibility: hidden;}
      [data-testid="stDecoration"] {display:none;}
      .viewerBadge_container__E0v7 {display:none !important;}
      /* Tighter selectboxes */
      div[data-baseweb="select"] > div {min-height: 34px;}
      label[for] {margin-bottom: 2px;}
    </style>
    """,
    unsafe_allow_html=True,
)

# ---------- deps ----------
try:
    import streamlit_authenticator as stauth
except Exception:
    st.error("streamlit-authenticator not installed. Add it to requirements.txt")
    st.stop()

# NOTE: python-docx is imported lazily inside to_docx_bytes() for faster cold start

# ---------- constants ----------
LOCAL_XLSX_DEFAULT = "Workorders.xlsx"

SHEET_WORKORDERS         = "Workorders"                  # history sheet
SHEET_ASSET_MASTER       = "Asset_Master"
SHEET_WO_MASTER          = "Workorders_Master"           # listing sheet with flags

# Service History sheet names (your current is Workorders_Master_Services)
SHEET_WO_SERVICE_CANDS   = [
    "Workorders_Master_Services",
    "Workorders_Master_service",
    "Workorders_Master_Service",
    "Workorders Service",
    "Service History"
]

# Service Report sheet candidate names
SHEET_SERVICE_CANDIDATES = ["Service Report", "Service_Report", "ServiceReport"]

# Optional Users sheet
SHEET_USERS_CANDIDATES   = ["Users", "Users]", "USERS", "users"]

REQUIRED_WO_COLS = [
    "WORKORDER","TITLE","STATUS","PO","P/N","QUANTITY RECEIVED",
    "Vendors","COMPLETED ON","ASSET","Location",
]
OPTIONAL_SORT_COL = "Sort"
ASSET_MASTER_COLS = ["Location","ASSET"]

MASTER_REQUIRED = [
    "ID","Title","Description","Asset","Status","Created on","Planned Start Date",
    "Due date","Started on","Completed on","Assigned to","Teams Assigned to",
    "Completed by","Location","IsOpen","IsOverdue","IsScheduled","IsCompleted","IsOld"
]

# Canon for Service Report
SERVICE_REPORT_CANON = {
    "WO_ID":{"workorder","wo","work order","work order id","id","wo id"},
    "Title":{"title","name"},
    "Service":{"service","procedure","procedure name","task","step","line item"},
    "Asset":{"asset","asset name"},
    "Location":{"location","ns location","location2"},
    "Date":{"date","completed on","performed on","service date","closed on"},
    "User":{"user","technician","completed by","performed by","assigned to"},
    "Notes":{"notes","description","comment","comments","details"},
    "Status":{"status"},
    "Schedule":{"schedule","interval","frequency","meter interval","planned interval","cycle"},
    "Remaining":{"remaining","remaining value","units remaining","miles remaining","hours remaining","reading remaining","remaining units"},
    "Percent Remaining":{"percent remaining","% remaining","remaining %","remaining pct","pct remaining"},
    "Meter Type":{"meter type","type","uom","unit","units"},
    "Due Date":{"due date","next due","target date","next service date"},
    # If your Service Report ever includes MReading, this will map it
    "MReading":{"mreading","meter reading","reading at service","reading"},
}

# Canon for Service History (Workorders_Master_Services)
SERVICE_HISTORY_CANON = {
    "WO_ID":{"id","wo","workorder","work order","workorder id"},
    "Title":{"title"},
    "Service":{"service type","service","procedure name","procedure","task"},
    "Asset":{"asset","asset name"},
    "Location2":{"location2","location","ns location"},
    "Date":{"completed on","performed on","date","service date"},
    "User":{"completed by","technician","assigned to","performed by","user"},
    "Notes":{"notes","description","comment","comments","details"},
    "Status":{"status"},
    "MReading":{"mreading","meter reading","reading at service","reading"},
    "MHours":{"mhours","hours at service","hours"},
}

# ---------- helpers ----------
def to_plain(obj):
    if isinstance(obj, Mapping):
        return {k: to_plain(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [to_plain(x) for x in obj]
    return obj

def load_config() -> dict:
    if "app_config" in st.secrets:
        return to_plain(st.secrets["app_config"])
    if "app_config_yaml" in st.secrets:
        try:
            return yaml.safe_load(st.secrets["app_config_yaml"]) or {}
        except Exception as e:
            st.error(f"Invalid YAML in app_config_yaml secret: {e}")
            return {}
    here = Path(__file__).resolve().parent
    cfg_file = here / "app_config.yaml"
    if cfg_file.exists():
        try:
            return yaml.safe_load(cfg_file.read_text(encoding="utf-8")) or {}
        except Exception as e:
            st.error(f"Invalid YAML in app_config.yaml: {e}")
            return {}
    return {}

def _norm_date_any(s: str) -> str:
    s = (str(s) if s is not None else "").strip()
    if not s:
        return ""
    for fmt in ("%Y-%m-%d","%m/%d/%Y","%m/%d/%y","%d-%b-%Y","%Y-%m-%d %H:%M:%S"):
        try:
            return datetime.strptime(s, fmt).strftime("%Y-%m-%d")
        except Exception:
            pass
    try:
        dt = pd.to_datetime(s, errors="coerce")
        return "" if pd.isna(dt) else dt.strftime("%Y-%m-%d")
    except Exception:
        return s

def _norm_key(x: str) -> str:
    s = re.sub(r"[^0-9a-z]+", " ", str(x).lower())
    return re.sub(r"\s+", " ", s).strip()

def _canonize_headers(df: pd.DataFrame, canon: dict[str, set[str]]) -> pd.DataFrame:
    low_to_orig = {str(c).strip().lower(): str(c) for c in df.columns}
    mapping = {}
    for key, aliases in canon.items():
        key_l = key.lower()
        if key_l in low_to_orig:
            mapping[low_to_orig[key_l]] = key
            continue
        for low, orig in low_to_orig.items():
            low2 = re.sub(r"\s+", " ", low)
            if low in aliases or low2 in aliases:
                mapping[orig] = key
                break
    return df.rename(columns=mapping)

def to_xlsx_bytes(df: pd.DataFrame, sheet: str) -> bytes:
    import xlsxwriter  # ensure dependency exists; used only on download
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="xlsxwriter") as w:
        df.to_excel(w, index=False, sheet_name=sheet)
        ws = w.sheets[sheet]
        ws.autofilter(0, 0, max(0, len(df)), max(0, len(df.columns)-1))
        for i, col in enumerate(df.columns):
            width = 12 if df.empty else min(60, max(10, int(df[col].astype(str).str.len().quantile(0.9)) + 2))
            ws.set_column(i, i, width)
    return buf.getvalue()

def to_docx_bytes(df: pd.DataFrame, title: str) -> bytes:
    # lazy import for faster cold start
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)
    doc.add_heading(title, level=1)
    rows, cols = len(df)+1, len(df.columns)
    tbl = doc.add_table(rows=rows, cols=cols); tbl.style = "Table Grid"
    for j, c in enumerate(df.columns): tbl.cell(0, j).text = str(c)
    for i, (_, r) in enumerate(df.iterrows(), start=1):
        for j, c in enumerate(df.columns):
            v = "" if pd.isna(r[c]) else str(r[c]); tbl.cell(i, j).text = v
    out = io.BytesIO(); doc.save(out); return out.getvalue()

def coerce_bool(s: pd.Series) -> pd.Series:
    if s.dtype == bool: return s
    m = s.astype(str).str.strip().str.lower()
    true_vals  = {"true","yes","y","1","t"}
    false_vals = {"false","no","n","0","f","", "nan", "none"}
    out = m.map(lambda x: True if x in true_vals else (False if x in false_vals else False))
    return out.astype(bool)

# --- Cross-Reference (local file) ---
XREF_LOCAL_DEFAULT = "Parts crossreference.xlsm"  # falls back to .xlsx if needed

def _xref_path_from_cfg(cfg: dict) -> Path:
    # lets you override via app_config.yaml -> settings.xref_path
    return Path((cfg.get("settings", {}) or {}).get("xref_path") or XREF_LOCAL_DEFAULT)

@st.cache_data(ttl=180, show_spinner=False)
def _load_xref_bytes(cfg: dict) -> bytes:
    p = _xref_path_from_cfg(cfg)
    if not p.exists():
        alt = p.with_suffix(".xlsx")
        if alt.exists():
            p = alt
        else:
            raise FileNotFoundError(f"Cross-reference workbook not found: {p.resolve()}")
    return p.read_bytes()

# --- Cross-Reference (local) — constants & helpers (no GitHub) ---
XREF_SHEET_INV   = "Parts_Master"
XREF_SHEET_MERGE = "Merge"

# Canonical inventory columns
XREF_INV_COL_PN    = "Part Numbers"
XREF_INV_COL_QTY   = "Quantity in Stock"
XREF_INV_COL_LOC2  = "Location2"
XREF_INV_COL_LOC   = "Location"
XREF_INV_COL_AREA  = "Area"
XREF_INV_COL_NAME  = "INV_NAME"  # standardized Name/Description

XREF_NAME_CANDIDATES = [
    "Name", "Description", "Item", "Part Name", "Description 1",
    "Item Description", "Product Name"
]

# Optional: prioritize brands when the same token appears under multiple columns
XREF_BRAND_PRIORITY = ["Caterpillar", "Donaldson", "Fleetguard", "Baldwin", "Wix", "Fram"]

# ----- normalization helpers -----
_XREF_NBSP = "\u00A0"

def xref_normalize_pn(pn) -> str:
    if pd.isna(pn): return ""
    s = str(pn).upper().replace(_XREF_NBSP, " ")
    return "".join(ch for ch in s if ch.isalnum())

def xref_last_token_with_digits(s: str) -> str:
    if not isinstance(s, str): s = str(s)
    tokens = re.findall(r"[A-Za-z0-9]+", s.upper())
    tokens = [t for t in tokens if any(ch.isdigit() for ch in t)]
    return tokens[-1] if tokens else ""

def xref_digits_only(s: str) -> str:
    return "".join(ch for ch in str(s) if ch.isdigit())

def xref_norm_text(s: str) -> str:
    if s is None: return ""
    s = str(s).upper().strip().replace(_XREF_NBSP, " ")
    s = re.sub(r"[^A-Z0-9 ]+", " ", s)
    s = re.sub(r"\s+", " ", s)
    return s

# ----- Excel open (supports .xlsm/.xlsx) -----
def xref_excel_from_bytes(b: bytes) -> pd.ExcelFile:
    return pd.ExcelFile(io.BytesIO(b), engine="openpyxl")

# ----- Inventory sheet loader -----
def xref_find_header_row_for_inventory(df_raw: pd.DataFrame):
    candidates = {"part numbers","quantity in stock","location2","location","area","qty","on hand","name","description"}
    for i in range(min(50, len(df_raw))):
        vals = [str(v).strip().lower() for v in df_raw.iloc[i].tolist()]
        if sum(v in candidates for v in vals) >= 2:
            headers, seen = [], {}
            for v in df_raw.iloc[i].tolist():
                name = str(v).strip() or "Unnamed"
                if name in seen:
                    seen[name]+=1; name=f"{name}.{seen[name]}"
                else:
                    seen[name]=0
                headers.append(name)
            return i, headers
    return None, None

def xref_finalize_inventory_dataframe(df_raw: pd.DataFrame) -> pd.DataFrame:
    header_idx, header_vals = xref_find_header_row_for_inventory(df_raw)
    if header_idx is None:
        df = df_raw.copy()
        df.columns = [str(c).strip() for c in df.iloc[0].tolist()]
        df = df.iloc[1:].reset_index(drop=True)
    else:
        df = df_raw.iloc[header_idx+1:].copy()
        df.columns = [str(c).strip() for c in header_vals]
        df = df.reset_index(drop=True)

    cols_clean = pd.Index(df.columns).astype(str).str.strip()
    df = df.loc[:, ~(cols_clean == "")]

    # rename to canonical
    rename_map = {}
    for col in df.columns:
        key = str(col).strip().lower()
        if key in {"part numbers","part number","pn"}: rename_map[col]=XREF_INV_COL_PN
        elif key in {"location2","location 2","loc2"}: rename_map[col]=XREF_INV_COL_LOC2
        elif key in {"quantity in stock","qty in stock","quantity","qty","on hand","qoh"}: rename_map[col]=XREF_INV_COL_QTY
        elif key in {"location","loc"}: rename_map[col]=XREF_INV_COL_LOC
        elif key == "area": rename_map[col]=XREF_INV_COL_AREA
    df = df.rename(columns=rename_map)

    if XREF_INV_COL_QTY in df.columns:
        df[XREF_INV_COL_QTY] = pd.to_numeric(df[XREF_INV_COL_QTY], errors="coerce").fillna(0)
    for c in df.columns:
        if df[c].dtype == object:
            df[c] = df[c].astype(str).str.strip()

    # Name/Description -> INV_NAME
    name_col = next((c for c in XREF_NAME_CANDIDATES if c in df.columns), None)
    df[XREF_INV_COL_NAME] = df[name_col].astype(str) if name_col else ""

    # Normalized keys for matching
    if XREF_INV_COL_PN in df.columns:
        df["_PN_NORM"] = df[XREF_INV_COL_PN].map(xref_normalize_pn)
        df["_PN_TAIL"] = df[XREF_INV_COL_PN].map(xref_last_token_with_digits).map(xref_normalize_pn)
        df["_PN_DIG"]  = df[XREF_INV_COL_PN].map(xref_digits_only)
    else:
        df["_PN_NORM"]=df["_PN_TAIL"]=df["_PN_DIG"]=""

    df["_NAME_NORM"] = df[XREF_INV_COL_NAME].map(xref_normalize_pn)
    df["_NAME_TAIL"] = df[XREF_INV_COL_NAME].map(xref_last_token_with_digits).map(xref_normalize_pn)
    df["_NAME_DIG"]  = df[XREF_INV_COL_NAME].map(xref_digits_only)

    # Stable row id (so per-row aggregation preserves location dupes)
    df["_ROWID"] = range(len(df))

    def stock_text(row) -> str:
        parts=[]
        if XREF_INV_COL_LOC2 in row and str(row[XREF_INV_COL_LOC2]).strip(): parts.append(str(row[XREF_INV_COL_LOC2]).strip())
        if XREF_INV_COL_LOC  in row and str(row[XREF_INV_COL_LOC]).strip():  parts.append(str(row[XREF_INV_COL_LOC]).strip())
        if XREF_INV_COL_AREA in row and str(row[XREF_INV_COL_AREA]).strip(): parts.append(str(row[XREF_INV_COL_AREA]).strip())
        if XREF_INV_COL_QTY  in row and str(row[XREF_INV_COL_QTY]).strip()!="":
            try: q = int(float(row[XREF_INV_COL_QTY]))
            except Exception: q = row[XREF_INV_COL_QTY]
            parts.append(f"Qty - {q}")
        return "; ".join(parts) if parts else ""
    df["_STOCK_TXT"] = df.apply(stock_text, axis=1)
    return df

@st.cache_data(ttl=180, show_spinner=False)
def xref_load_inventory_df(cfg: dict, debug=False) -> pd.DataFrame:
    xls = xref_excel_from_bytes(_load_xref_bytes(cfg))
    df_raw = pd.read_excel(xls, sheet_name=XREF_SHEET_INV, header=None)
    df = xref_finalize_inventory_dataframe(df_raw)
    if debug:
        st.sidebar.write("**XRef: Detected inventory columns:**")
        for c in df.columns: st.sidebar.write(f"- {c}")
    return df

# ----- Merge sheet loader -----
def _xref_find_header_row_for_merge(df_raw: pd.DataFrame) -> int:
    def is_brandish(v: str) -> bool:
        v = str(v).strip()
        if v == "" or v.lower() == "nan": return False
        letters = sum(ch.isalpha() for ch in v)
        digits  = sum(ch.isdigit() for ch in v)
        return letters >= 2 and letters >= digits and len(v) <= 30
    for i in range(min(50, len(df_raw))):
        row = df_raw.iloc[i].tolist()
        nonblank = [x for x in row if str(x).strip() not in {"", "nan", "None"}]
        brandish = sum(is_brandish(x) for x in row)
        if len(nonblank) >= 3 and brandish >= 2: return i
    return 0

_XREF_SPLIT_RE = re.compile(r"[;,/\n]+")

@st.cache_data(ttl=180, show_spinner=False)
def xref_load_merge_wide_and_long(cfg: dict, debug=False):
    xls = xref_excel_from_bytes(_load_xref_bytes(cfg))
    df_raw = pd.read_excel(xls, sheet_name=XREF_SHEET_MERGE, header=None)

    hdr_idx = _xref_find_header_row_for_merge(df_raw)
    raw_headers = df_raw.iloc[hdr_idx].tolist()
    headers, seen = [], {}
    for j, v in enumerate(raw_headers):
        name = str(v).strip() or f"Col_{j}"
        if name in seen:
            seen[name]+=1; name=f"{name}.{seen[name]}"
        else:
            seen[name]=0
        headers.append(name)

    merge_df = df_raw.iloc[hdr_idx+1:].copy()
    merge_df.columns = headers
    merge_df = merge_df.reset_index(drop=True)

    for c in merge_df.columns:
        if merge_df[c].dtype == object:
            merge_df[c] = merge_df[c].astype(str).str.strip()
    keep = merge_df.columns[merge_df.apply(lambda s: s.astype(str).str.strip().ne("").any())]
    merge_df = merge_df[keep]

    if debug:
        st.sidebar.write("**XRef: Detected MERGE brands:**")
        for c in merge_df.columns: st.sidebar.write(f"- {c}")

    wide = merge_df.reset_index().rename(columns={"index":"RowID"})
    parts = []
    for _, row in wide.iterrows():
        rid = int(row["RowID"])
        for brand in wide.columns:
            if brand == "RowID": continue
            raw_val = str(row[brand]).strip()
            if raw_val and raw_val.lower() != "nan":
                for piece in _XREF_SPLIT_RE.split(raw_val):
                    piece = piece.strip()
                    if piece:
                        parts.append((rid, brand, piece))
    long = pd.DataFrame(parts, columns=["RowID","Brand","PartNumber"])
    if long.empty:
        long = pd.DataFrame(columns=["RowID","Brand","PartNumber"])

    # normalized match keys
    long["PN_Full"]   = long["PartNumber"].map(xref_normalize_pn)
    long["PN_Tail"]   = long["PartNumber"].map(xref_last_token_with_digits).map(xref_normalize_pn)
    long["PN_Digits"] = long["PartNumber"].map(xref_digits_only)

    return merge_df, long

# ----- Resolve + build + inventory match -----
def _xref_apply_brand_priority(df_hits: pd.DataFrame) -> pd.DataFrame:
    if df_hits.empty or not XREF_BRAND_PRIORITY:
        return df_hits
    prio = {b.upper(): i for i, b in enumerate(XREF_BRAND_PRIORITY)}
    return df_hits.assign(_p=df_hits["Brand"].str.upper().map(prio).fillna(999)) \
                  .sort_values(by=["_p","PartNumber","RowID"]).drop(columns="_p")

def xref_resolve_rowset_and_primary_brand(query: str, long_df: pd.DataFrame):
    q_full   = xref_normalize_pn(query)
    q_tail   = xref_normalize_pn(xref_last_token_with_digits(query))
    q_digits = xref_digits_only(query)

    buckets = []
    buckets.append(long_df[long_df["PN_Full"] == q_full])
    if q_tail and len(q_tail) >= 5:
        buckets.append(long_df[long_df["PN_Tail"] == q_tail])
    if q_digits and len(q_digits) >= 5:
        buckets.append(long_df[long_df["PN_Digits"] == q_digits])

    for cand in buckets:
        if not cand.empty:
            row_ids = set(cand["RowID"].astype(int).tolist())
            c = _xref_apply_brand_priority(cand).assign(_len=cand["PartNumber"].astype(str).str.len())
            best = c.sort_values(by=["_len","RowID"]).iloc[0]
            return row_ids, str(best["Brand"])
    return set(), None

def xref_build_crossrefs_from_rows(merge_df: pd.DataFrame, row_ids: set) -> pd.DataFrame:
    if not row_ids:
        return pd.DataFrame(columns=["Brand","PartNumber","PN_Norm"])
    frames = []
    for rid in sorted(row_ids):
        row = merge_df.loc[rid]
        items = []
        for brand in merge_df.columns:
            val = str(row[brand]).strip()
            if val and val.lower()!="nan":
                items.append({"Brand":brand,"PartNumber":val,"PN_Norm":xref_normalize_pn(val)})
        if items:
            frames.append(pd.DataFrame(items))
    xrefs = pd.concat(frames, ignore_index=True) if frames else pd.DataFrame(columns=["Brand","PartNumber","PN_Norm"])
    if xrefs.empty:
        return xrefs
    prio = {b.upper(): i for i, b in enumerate(XREF_BRAND_PRIORITY)}
    xrefs["_p"] = xrefs["Brand"].str.upper().map(prio).fillna(999)
    xrefs = (xrefs.sort_values(by=["PN_Norm","_p","PartNumber"])
                   .drop_duplicates(subset=["PN_Norm"], keep="first")
                   .drop(columns="_p")
                   .reset_index(drop=True))
    return xrefs

def xref_add_entered_pn_to_xrefs(xrefs: pd.DataFrame, pn_input: str, matched_brand: str|None) -> pd.DataFrame:
    pn_norm = xref_normalize_pn(pn_input)
    if pn_norm and pn_norm not in set(xrefs["PN_Norm"]):
        brand = matched_brand if matched_brand else "Entered PN"
        extra = pd.DataFrame([{"Brand":brand,"PartNumber":pn_input,"PN_Norm":pn_norm}])
        xrefs = pd.concat([extra, xrefs], ignore_index=True)
    return xrefs

def xref_inventory_lookup_per_xref(inv_df: pd.DataFrame, cross_df: pd.DataFrame, location2: str|None) -> pd.DataFrame:
    df = inv_df.copy()

    # Filters: Location2 + in-stock only
    if location2 and XREF_INV_COL_LOC2 in df.columns:
        loc_norm = xref_norm_text(location2)
        df = df[df[XREF_INV_COL_LOC2].map(xref_norm_text) == loc_norm]
    if XREF_INV_COL_QTY in df.columns:
        df[XREF_INV_COL_QTY] = pd.to_numeric(df[XREF_INV_COL_QTY], errors="coerce").fillna(0)
        df = df[df[XREF_INV_COL_QTY] > 0]

    def S(col: str) -> pd.Series:
        return df[col] if col in df.columns else pd.Series([""]*len(df), index=df.index)

    hits_all = []

    for _, xr in cross_df.iterrows():
        brand = str(xr.get("Brand","")).strip()
        pn    = str(xr.get("PartNumber","")).strip()
        if not pn:
            continue

        pn_norm   = xref_normalize_pn(pn)
        tail_norm = xref_normalize_pn(xref_last_token_with_digits(pn))
        dig       = xref_digits_only(pn)
        brand_pn  = xref_normalize_pn(f"{brand} {pn}") if brand else ""

        m = pd.Series(False, index=df.index)

        # Exact on PN / Name
        m = m | (S("_PN_NORM")  == pn_norm) | (S("_NAME_NORM")  == pn_norm)

        # Exact on "Brand + PN"
        if brand_pn:
            m = m | (S("_PN_NORM") == brand_pn) | (S("_NAME_NORM") == brand_pn)

        # Exact tail + ends-with tail
        if tail_norm and len(tail_norm) >= 5:
            m = m | (S("_PN_TAIL") == tail_norm) | (S("_NAME_TAIL") == tail_norm)
            m = m | S("_PN_NORM").astype(str).str.endswith(tail_norm, na=False)
            m = m | S("_NAME_NORM").astype(str).str.endswith(tail_norm, na=False)

        # Exact digits + ends-with digits
        if len(dig) >= 5:
            m = m | (S("_PN_DIG") == dig) | (S("_NAME_DIG") == dig)
            m = m | S("_PN_DIG").astype(str).str.endswith(dig, na=False)
            m = m | S("_NAME_DIG").astype(str).str.endswith(dig, na=False)

        matched = df.loc[m].copy()
        if matched.empty:
            continue

        # Prepare output (preserve location detail)
        show_cols = []
        if XREF_INV_COL_PN in matched.columns:   show_cols.append(XREF_INV_COL_PN)
        if XREF_INV_COL_NAME in matched.columns: show_cols.append(XREF_INV_COL_NAME)
        for c in (XREF_INV_COL_LOC2, XREF_INV_COL_LOC, XREF_INV_COL_AREA, XREF_INV_COL_QTY):
            if c in matched.columns: show_cols.append(c)
        show_cols += ["_STOCK_TXT","_ROWID"]
        show_cols = [c for c in show_cols if c in matched.columns]

        out = matched[show_cols].rename(columns={
            XREF_INV_COL_PN:   "Inventory PN",
            XREF_INV_COL_NAME: "Name",
            XREF_INV_COL_QTY:  "Qty",
            "_STOCK_TXT":      "Stock"
        })
        out.insert(0, "Matched PN", pn)
        out.insert(0, "Matched Brand", brand)
        hits_all.append(out)

    if not hits_all:
        return pd.DataFrame(columns=["Matched From","Inventory PN","Name","Location2","Location","Area","Qty","Stock"])

    result = pd.concat(hits_all, ignore_index=True)
    result["Matched Pair"] = (result.get("Matched Brand","").astype(str) + " " + result.get("Matched PN","").astype(str)).str.strip()

    # Aggregate **per inventory row** (_ROWID) so duplicates by location remain separate
    group_keys = ["_ROWID"]
    keep_cols = [c for c in ["Inventory PN","Name","Location2","Location","Area","Qty","Stock"] if c in result.columns]
    agg = result.groupby(group_keys, as_index=False).agg(**{
        "Matched From": ("Matched Pair", lambda s: ", ".join(sorted({x for x in s if x})))
    })
    static = result.drop_duplicates("_ROWID")[["_ROWID"] + keep_cols]
    merged = pd.merge(agg, static, on="_ROWID", how="left").drop(columns=["_ROWID"])

    order = [c for c in ["Matched From","Inventory PN","Name","Location2","Location","Area","Qty","Stock"] if c in merged.columns]
    merged = merged[order]

    sort_cols = [c for c in ["Inventory PN","Name","Location2","Location","Area"] if c in merged.columns]
    return merged.sort_values(by=sort_cols).reset_index(drop=True)

# ---------- fast I/O / Parquet layer ----------
PARQUET_DIR = Path("faststore")
PARQUET_DIR.mkdir(exist_ok=True)

def _xlsx_path_from_cfg(cfg: dict) -> Path:
    return Path((cfg.get("settings", {}) or {}).get("xlsx_path") or LOCAL_XLSX_DEFAULT)

def _mtime(path: Path) -> float:
    try:
        return path.stat().st_mtime
    except Exception:
        return 0.0

@st.cache_data(show_spinner=False)
def get_local_xlsx_bytes_cached(xlsx_path_str: str, xlsx_mtime: float) -> bytes:
    # cache keyed by path + mtime so reruns don't re-read file from disk
    return Path(xlsx_path_str).read_bytes()

def _pq_path(sheet: str) -> Path:
    safe = re.sub(r"[^0-9A-Za-z_.-]+", "_", sheet.strip())
    return PARQUET_DIR / f"{safe}.parquet"

def _write_parquet(df: pd.DataFrame, sheet: str) -> None:
    df.to_parquet(_pq_path(sheet), index=False)

def _parquet_fresh(sheet: str, xlsx_path: Path) -> bool:
    pq = _pq_path(sheet)
    return pq.exists() and _mtime(pq) >= _mtime(xlsx_path)

def _read_parquet_columns(sheet: str, columns: list[str] | None = None) -> pd.DataFrame:
    return pd.read_parquet(_pq_path(sheet), columns=columns)

def _rebuild_parquet_from_excel(xlsx_bytes: bytes, sheets_to_pull: list[str]) -> dict[str, pd.DataFrame]:
    # Parse multiple sheets with one Excel decode, then write parquet
    with pd.ExcelFile(io.BytesIO(xlsx_bytes), engine="openpyxl") as xf:
        out = {}
        for sh in sheets_to_pull:
            try:
                df = xf.parse(sh, dtype=str, keep_default_na=False)
                df.columns = [str(c).strip() for c in df.columns]
                _write_parquet(df, sh)
                out[sh] = df
            except Exception:
                pass
    return out

def parquet_button_and_refresh(xlsx_path: Path, xlsx_bytes: bytes) -> None:
    with st.sidebar.expander("⚡ Data cache"):
        if st.button("Rebuild Parquet cache now"):
            _rebuild_parquet_from_excel(xlsx_bytes, [
                SHEET_WORKORDERS, SHEET_ASSET_MASTER, SHEET_WO_MASTER,
                *SHEET_WO_SERVICE_CANDS, *SHEET_SERVICE_CANDIDATES, *SHEET_USERS_CANDIDATES
            ])
            st.success("Parquet cache rebuilt.")
            st.cache_data.clear()
            st.rerun()

# ---------- data access ----------
def get_local_xlsx_bytes(cfg: dict) -> bytes:
    p = _xlsx_path_from_cfg(cfg)
    if not p.exists():
        raise FileNotFoundError(f"Local Excel not found: {p.resolve()}")
    return get_local_xlsx_bytes_cached(str(p), _mtime(p))

def get_data_last_updated_local(cfg: dict) -> str | None:
    xl = (cfg.get("settings", {}) or {}).get("xlsx_path") or LOCAL_XLSX_DEFAULT
    try:
        from zoneinfo import ZoneInfo
        ts = datetime.fromtimestamp(os.path.getmtime(xl), tz=timezone.utc).astimezone(ZoneInfo("America/New_York"))
        return ts.strftime("Data last updated: %Y-%m-%d %H:%M ET")
    except Exception:
        return None

# ---------- loaders (Parquet fast-path) ----------
@st.cache_data(show_spinner=False)
def load_workorders_df(xlsx_mtime: float, xlsx_path_str: str) -> pd.DataFrame:
    need = REQUIRED_WO_COLS + ([OPTIONAL_SORT_COL] if OPTIONAL_SORT_COL else [])
    sheet = SHEET_WORKORDERS
    xlsx_path = Path(xlsx_path_str)

    if _parquet_fresh(sheet, xlsx_path):
        df = _read_parquet_columns(sheet, columns=[c for c in need if c])
    else:
        _rebuild_parquet_from_excel(get_local_xlsx_bytes_cached(xlsx_path_str, xlsx_mtime), [sheet])
        df = _read_parquet_columns(sheet, columns=[c for c in need if c])

    df.columns = [str(c).strip() for c in df.columns]
    missing = [c for c in REQUIRED_WO_COLS if c not in df.columns]
    if missing:
        raise ValueError(f"Sheet '{sheet}' missing columns: {missing}\nFound: {list(df.columns)}")

    df = df[[c for c in need if c in df.columns]].copy()
    if "COMPLETED ON" in df.columns:
        df["COMPLETED ON"] = df["COMPLETED ON"].map(_norm_date_any)
    for c in df.columns:
        if df[c].dtype == object:
            df[c] = df[c].astype("string").str.strip()
    for catcol in ("Location","ASSET"):
        if catcol in df.columns:
            df[catcol] = df[catcol].astype("category")
    return df

@st.cache_data(show_spinner=False)
def load_asset_master_df(xlsx_mtime: float, xlsx_path_str: str) -> pd.DataFrame:
    sheet = SHEET_ASSET_MASTER
    xlsx_path = Path(xlsx_path_str)
    if _parquet_fresh(sheet, xlsx_path):
        df = _read_parquet_columns(sheet, columns=ASSET_MASTER_COLS)
    else:
        _rebuild_parquet_from_excel(get_local_xlsx_bytes_cached(xlsx_path_str, xlsx_mtime), [sheet])
        df = _read_parquet_columns(sheet, columns=ASSET_MASTER_COLS)

    df.columns = [str(c).strip() for c in df.columns]
    for c in ASSET_MASTER_COLS:
        if c not in df.columns:
            raise ValueError(f"Sheet '{sheet}' missing '{c}'")
        df[c] = df[c].astype("string").str.strip()
    df = df[(df["Location"] != "") & (df["ASSET"] != "")]
    df["Location"] = df["Location"].astype("category")
    df["ASSET"] = df["ASSET"].astype("category")
    return df[ASSET_MASTER_COLS].copy()

@st.cache_data(show_spinner=False)
def load_wo_master_df(xlsx_mtime: float, xlsx_path_str: str) -> pd.DataFrame:
    sheet = SHEET_WO_MASTER
    xlsx_path = Path(xlsx_path_str)
    if _parquet_fresh(sheet, xlsx_path):
        df = _read_parquet_columns(sheet, columns=None)
    else:
        _rebuild_parquet_from_excel(get_local_xlsx_bytes_cached(xlsx_path_str, xlsx_mtime), [sheet])
        df = _read_parquet_columns(sheet, columns=None)

    df.columns = [str(c).strip() for c in df.columns]
    have = [c for c in MASTER_REQUIRED if c in df.columns]
    if have:
        df = df[have].copy()

    for dc in ("Created on","Planned Start Date","Due date","Started on","Completed on"):
        if dc in df.columns:
            df[dc] = df[dc].map(_norm_date_any)
    for bc in ("IsOpen","IsOverdue","IsScheduled","IsCompleted","IsOld"):
        if bc in df.columns:
            df[bc] = coerce_bool(df[bc])
    for c in [x for x in ["ID","Title","Description","Asset","Status","Assigned to","Teams Assigned to","Completed by","Location"] if x in df.columns]:
        df[c] = df[c].astype("string").str.strip()
    if "ID" in df.columns:
        df["ID"] = df["ID"].astype("string").str.strip()

    for catcol in ("Location","Assigned to","Asset"):
        if catcol in df.columns:
            df[catcol] = df[catcol].astype("category")
    return df

@st.cache_data(show_spinner=False)
def load_service_report_df(xlsx_mtime: float, xlsx_path_str: str):
    xlsx_path = Path(xlsx_path_str)

    for nm in SHEET_SERVICE_CANDIDATES:
        if _parquet_fresh(nm, xlsx_path):
            raw = _read_parquet_columns(nm, columns=None)
        else:
            try:
                _rebuild_parquet_from_excel(get_local_xlsx_bytes_cached(xlsx_path_str, xlsx_mtime), [nm])
                if _parquet_fresh(nm, xlsx_path):
                    raw = _read_parquet_columns(nm, columns=None)
                else:
                    continue
            except Exception:
                continue

        raw.columns = [str(c).strip() for c in raw.columns]
        canon = _canonize_headers(raw.copy(), SERVICE_REPORT_CANON)
        if "Date" in canon.columns: canon["Date"] = canon["Date"].map(_norm_date_any)
        if "Due Date" in canon.columns: canon["Due Date"] = canon["Due Date"].map(_norm_date_any)

        for col, newcol in [("Schedule","__Schedule_num"), ("Remaining","__Remaining_num"), ("Percent Remaining","__PctRemain_num")]:
            if col in canon.columns:
                canon[newcol] = pd.to_numeric(canon[col].astype(str).str.replace("%","", regex=False), errors="coerce")
            else:
                canon[newcol] = pd.NA
        if "__PctRemain_num" in canon.columns:
            pr = pd.to_numeric(canon["__PctRemain_num"], errors="coerce")
            canon["__PctRemain_num"] = pr.where((pr.isna()) | (pr <= 1.0), pr/100.0)

        canon["__MeterType_norm"] = canon.get("Meter Type", pd.Series([], dtype=str)).astype(str).str.strip().str.lower() if "Meter Type" in canon.columns else ""
        canon["__Due_dt"] = pd.to_datetime(canon["Due Date"], errors="coerce") if "Due Date" in canon.columns else pd.NaT
        for c in [x for x in ["WO_ID","Title","Service","Asset","Location","User","Notes","Status","MReading"] if x in canon.columns]:
            canon[c] = canon[c].astype("string").str.strip()
        return raw, canon, nm

    return None, None, None

@st.cache_data(show_spinner=False)
def load_service_history_df(xlsx_mtime: float, xlsx_path_str: str):
    xlsx_path = Path(xlsx_path_str)
    last_err = None
    for nm in SHEET_WO_SERVICE_CANDS:
        try:
            if _parquet_fresh(nm, xlsx_path):
                df = _read_parquet_columns(nm, columns=None)
            else:
                _rebuild_parquet_from_excel(get_local_xlsx_bytes_cached(xlsx_path_str, xlsx_mtime), [nm])
                df = _read_parquet_columns(nm, columns=None)

            df.columns = [str(c).strip() for c in df.columns]
            df = _canonize_headers(df, SERVICE_HISTORY_CANON)
            if "Date" in df.columns: df["Date"] = df["Date"].map(_norm_date_any)
            for c in [x for x in ["WO_ID","Title","Service","Asset","Location2","User","Notes","Status","MReading","MHours"] if x in df.columns]:
                df[c] = df[c].astype("string").str.strip()
            keep = [c for c in ["Date","WO_ID","Title","Service","MReading","MHours","Asset","User","Location2","Notes","Status"] if c in df.columns]
            df = df[keep].copy() if keep else df
            return df, nm
        except Exception as e:
            last_err = e
            continue
    return None, f"{last_err}" if last_err else None

@st.cache_data(show_spinner=False)
def load_users_sheet(xlsx_mtime: float, xlsx_path_str: str) -> list[str] | None:
    xlsx_path = Path(xlsx_path_str)
    for name in SHEET_USERS_CANDIDATES:
        try:
            if _parquet_fresh(name, xlsx_path):
                dfu = _read_parquet_columns(name, columns=None)
            else:
                _rebuild_parquet_from_excel(get_local_xlsx_bytes_cached(xlsx_path_str, xlsx_mtime), [name])
                dfu = _read_parquet_columns(name, columns=None)
            cols_low = {c.lower(): c for c in dfu.columns}
            col = cols_low.get("user")
            if not col:
                continue
            users = [u.strip() for u in dfu[col].astype(str).tolist() if str(u).strip()]
            users = sorted(dict.fromkeys(users))
            return users
        except Exception:
            pass
    return None

# ---------- Service Procedure Filter (LOCAL workbook) — page function ----------
def render_service_procedure_page(xmtime: float, xlsx_path_str: str):
    st.markdown("### Service Procedure Filter")
    # ---------- Optional Word export ----------
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.section import WD_ORIENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        DOCX_AVAILABLE = True
    except Exception:
        DOCX_AVAILABLE = False

    SHOW_PRINT = (platform.system() == "Windows") and (os.environ.get("ALLOW_SERVER_PRINT") == "1")

    # ---------- Sheet / Column names ----------
    SHEET_PROC     = "Service Procedures"
    SHEET_CTRL     = "Controls"
    SHEET_INV      = "Parts_Master"
    SHEET_XREF     = "Filter_List"

    COL_ASSET    = "Asset"
    COL_SERIAL   = "Serial"
    COL_SERVICE  = "Service"
    COL_TASKNO   = "Task #"
    COL_TASK     = "Task"
    COL_LOC_LIST = "Locations"

    INV_COL_PN    = "Part Numbers"
    INV_COL_QTY   = "Quantity in Stock"
    INV_COL_LOC   = "Location"
    INV_COL_AREA  = "Area"
    INV_COL_LOC2  = "Location2"

    XREF_COL_CAT  = "Cat"
    XREF_COL_DON  = "Donaldson P/N"

    PART_PREFIXES = ("Part Number_", "Part Type_", "Description_", "Qty_")

    # Optional inventory name/description candidates (if present we’ll use them)
    INV_NAME_CANDIDATES = [
        "Name", "Description", "Item", "Part Name", "Description 1",
        "Item Description", "Product Name"
    ]

    # ---------- Cached local sheet reader ----------
    xlsx_path = Path(xlsx_path_str)
    def read_sheet_cached(sheet: str) -> pd.DataFrame:
        if _parquet_fresh(sheet, xlsx_path):
            df = _read_parquet_columns(sheet, columns=None)
        else:
            _rebuild_parquet_from_excel(get_local_xlsx_bytes_cached(xlsx_path_str, xmtime), [sheet])
            df = _read_parquet_columns(sheet, columns=None)
        # de-duplicate any repeated column names
        return df.loc[:, ~pd.Index(df.columns).duplicated(keep="first")]

    # ---------- Load workbook (LOCAL) ----------
    try:
        df_proc = read_sheet_cached(SHEET_PROC)
        df_ctrl = read_sheet_cached(SHEET_CTRL)
        df_inv  = read_sheet_cached(SHEET_INV)
        df_xref = read_sheet_cached(SHEET_XREF) if _parquet_fresh(SHEET_XREF, xlsx_path) or True else pd.DataFrame()
    except Exception as e:
        st.error(f"Error loading workbook from local file: {e}")
        st.stop()

    with st.sidebar:
        st.caption(f"Workbook source: {xlsx_path.resolve()}")
        if st.button("Reload local workbook"):
            st.cache_data.clear()
            st.rerun()

    # ---------- Validate columns ----------
    need_proc = {COL_SERIAL, COL_SERVICE, COL_TASKNO, COL_TASK}
    need_ctrl = {COL_ASSET, COL_SERIAL, COL_LOC_LIST}
    need_inv  = {INV_COL_PN, INV_COL_QTY, INV_COL_LOC, INV_COL_AREA, INV_COL_LOC2}
    errors = []
    if not need_proc.issubset(df_proc.columns): errors.append("Service Procedures missing required columns.")
    if not need_ctrl.issubset(df_ctrl.columns): errors.append("Controls missing Asset/Serial/Locations.")
    if not need_inv.issubset(df_inv.columns):   errors.append("Parts_Master missing {Part Numbers, Quantity in Stock, Location, Area, Location2}.")
    if errors:
        st.error(" ".join(errors))
        st.stop()

    # ---------- Precompute inventory (Name-first, then Part Numbers) ----------
    df_inv = df_inv.copy()

    inv_name_col = next((c for c in INV_NAME_CANDIDATES if c in df_inv.columns), None)

    def _clean_space(s):
        if s is None:
            return ""
        # remove NBSPs/tabs and trim
        return str(s).replace("\u00A0", "").replace("\t", "").strip()

    _SPLIT_RE = re.compile(r"[,\;/\s]+")
    def normalize_pn(pn):
        """Uppercase, remove all non-alphanumerics (makes hyphens, spaces, dots irrelevant)."""
        if pd.isna(pn): return ""
        return "".join(ch for ch in str(pn).upper() if ch.isalnum())

    def _token_norms_from_text(s: str):
        """
        Produce PN-like tokens from free text by splitting on comma/semicolon/slash/space,
        then normalizing each token (hyphens/spaces removed), keeping only tokens that contain a digit.
        """
        if s is None: return []
        toks = []
        for piece in _SPLIT_RE.split(str(s).upper()):
            piece = piece.strip()
            if not piece:
                continue
            n = normalize_pn(piece)
            if n and any(ch.isdigit() for ch in n):
                toks.append(n)
        return toks

    def _semi_join(tokens):
        # Build ;T1;T2; string for exact token membership checks (Excel-like)
        uniq, seen = [], set()
        for t in tokens:
            if t not in seen:
                uniq.append(t); seen.add(t)
        return ";" + ";".join(uniq) + ";" if uniq else ";"

    def _build_tok_name(row) -> str:
        if not inv_name_col:
            return ";"
        toks = _token_norms_from_text(row.get(inv_name_col, ""))
        return _semi_join(toks)

    def _build_tok_pn(row) -> str:
        toks = _token_norms_from_text(row.get(INV_COL_PN, ""))
        return _semi_join(toks)

    # exact-token membership strings
    df_inv["_TOK_NAME_SEMI"] = df_inv.apply(_build_tok_name, axis=1) if inv_name_col else ";"
    df_inv["_TOK_PN_SEMI"]   = df_inv.apply(_build_tok_pn,   axis=1)

    # cleaned location + numeric qty
    df_inv["_LOC2_CLEAN"] = df_inv[INV_COL_LOC2].map(_clean_space)
    df_inv["_QTY_NUM"]    = pd.to_numeric(df_inv[INV_COL_QTY], errors="coerce").fillna(0)

    def excel_like_first_match(inv_df: pd.DataFrame, pn_norm: str, loc_val: str):
        """
        Prefer a token match in Name (if present), else fall back to Part Numbers.
        Matching is exact-token against prebuilt ;T1;T2; strings.
        """
        if pn_norm == "":
            return None

        loc_clean = _clean_space(loc_val)
        base = (inv_df["_LOC2_CLEAN"] == loc_clean) & (inv_df["_QTY_NUM"] > 0)

        # 1) Name-first match
        if "_TOK_NAME_SEMI" in inv_df.columns:
            m1 = inv_df["_TOK_NAME_SEMI"].str.contains(";" + pn_norm + ";", regex=False, na=False)
            sub1 = inv_df.loc[base & m1]
            if not sub1.empty:
                return sub1.iloc[0]

        # 2) Fallback: Part Numbers list
        m2 = inv_df["_TOK_PN_SEMI"].str.contains(";" + pn_norm + ";", regex=False, na=False)
        sub2 = inv_df.loc[base & m2]
        if not sub2.empty:
            return sub2.iloc[0]

        return None

    def inv_text_from_row(row):
        if row is None:
            return "No Stock"
        loc  = str(row.get(INV_COL_LOC, "")).strip()
        area = str(row.get(INV_COL_AREA, "")).strip()
        qty  = row.get(INV_COL_QTY)
        try:
            q = int(float(qty))
        except Exception:
            q = qty
        parts = []
        if loc:  parts.append(loc)
        if area: parts.append(area)
        if q is not None and str(q) != "nan":
            parts.append(f"Qty - {q}")
        return "; ".join(parts) if parts else "No Stock"

    # ---------- Unpivot service procedures to Task/Part rows ----------
    def unpivot_to_task_then_parts(df_proc_filtered: pd.DataFrame) -> pd.DataFrame:
        cols = list(df_proc_filtered.columns)
        part_cols = [c for c in cols if c.startswith(PART_PREFIXES)]

        # Header rows
        hdr = df_proc_filtered[[COL_TASKNO, COL_TASK]].drop_duplicates().copy()
        hdr["Part Number"] = None
        hdr["Part Type"]   = None
        hdr["Qty"]         = None
        hdr["RowKind"]     = "Task"
        hdr["SfxKey"]      = 0

        if not part_cols:
            out = hdr.copy()
            out["__TaskNum"] = pd.to_numeric(out[COL_TASKNO], errors="coerce").fillna(0)
            out = out.sort_values(by=["__TaskNum", "RowKind", "SfxKey"]).drop(columns="__TaskNum")
            return out[[COL_TASKNO, COL_TASK, "Part Number", "Part Type", "Qty"]].reset_index(drop=True)

        long = df_proc_filtered.melt(
            id_vars=[c for c in cols if c not in part_cols],
            value_vars=part_cols, var_name="Attr", value_name="Val"
        )
        long["Kind"] = long["Attr"].str.extract(r"^(.*)_", expand=False)
        long["Sfx"]  = long["Attr"].str.extract(r"_(.*)$",  expand=False)
        long = long.drop(columns=["Attr"])

        def first_non_null(x: pd.Series):
            x = x.dropna()
            return x.iloc[0] if not x.empty else None

        idx_cols = [c for c in [COL_TASKNO, COL_TASK, COL_SERVICE, COL_SERIAL, "Sfx"] if c in long.columns]
        wide = long.pivot_table(index=idx_cols, columns="Kind", values="Val",
                                aggfunc=first_non_null).reset_index()

        has_pt  = "Part Type" in wide.columns
        has_dsc = "Description" in wide.columns
        if has_pt and has_dsc:
            wide["Part Type (Display)"] = wide["Part Type"].where(
                wide["Part Type"].astype(str).str.strip().ne(""),
                wide["Description"]
            )
        elif has_pt:
            wide["Part Type (Display)"] = wide["Part Type"]
        elif has_dsc:
            wide["Part Type (Display)"] = wide["Description"]
        else:
            wide["Part Type (Display)"] = None

        if "Qty" in wide.columns:
            wide["Qty"] = pd.to_numeric(wide["Qty"], errors="coerce")

        sel_cols = [c for c in [COL_TASKNO, COL_TASK, "Part Number", "Part Type (Display)", "Qty", "Sfx"] if c in wide.columns]
        parts = wide[sel_cols].rename(columns={"Part Type (Display)": "Part Type"}).copy()
        parts["RowKind"] = "Part"
        parts["SfxKey"]  = pd.to_numeric(parts.get("Sfx"), errors="coerce").fillna(9999).astype(int)
        if "Sfx" in parts.columns:
            parts.drop(columns=["Sfx"], inplace=True)

        # Align and combine
        final_cols = [COL_TASKNO, COL_TASK, "Part Number", "Part Type", "Qty", "RowKind", "SfxKey"]

        def coerce_cols(df: pd.DataFrame) -> pd.DataFrame:
            df = df.loc[:, ~pd.Index(df.columns).duplicated(keep="first")]
            for c in final_cols:
                if c not in df.columns:
                    df[c] = None
            return df[final_cols]

        hdr   = coerce_cols(hdr)
        parts = coerce_cols(parts)

        combined = pd.concat([hdr, parts], ignore_index=True, sort=False)
        combined["__TaskNum"] = pd.to_numeric(combined[COL_TASKNO], errors="coerce").fillna(0)
        combined["__rk"]      = combined["RowKind"].map({"Task": 0, "Part": 1}).fillna(1)
        combined = combined.sort_values(by=["__TaskNum", "__rk", "SfxKey"]).drop(columns=["__TaskNum", "__rk"])

        combined.loc[combined["RowKind"] == "Part", COL_TASKNO] = None
        return combined[[COL_TASKNO, COL_TASK, "Part Number", "Part Type", "Qty"]].reset_index(drop=True)

    def _repeat_header(row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        trPr.append(tblHeader)

    def to_docx_bytes_sp(df: pd.DataFrame, *, asset: str, serial: str, service: str, location: str) -> bytes:
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        for side in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(section, side, Inches(0.5))

        doc.styles['Normal'].font.name = "Calibri"
        doc.styles['Normal'].font.size = Pt(10)

        title = doc.add_paragraph()
        run = title.add_run(
            f"Service Procedure Filter — Asset: {asset}  (Serial: {serial})  |  Service: {service}  |  Location: {location}"
        )
        run.bold = True
        doc.add_paragraph("")

        cols = list(df.columns)
        table = doc.add_table(rows=1, cols=len(cols))
        table.autofit = True
        hdr_cells = table.rows[0].cells
        for i, c in enumerate(cols):
            p = hdr_cells[i].paragraphs[0]
            r = p.add_run(str(c)); r.bold = True
            tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), "DDDDDD"); tcPr.append(shd)
        _repeat_header(table.rows[0])

        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, c in enumerate(cols):
                val = row.get(c, "")
                cells[i].text = "" if pd.isna(val) else str(val)

        width_map = {
            "Task #": 0.6, "Task": 1.8,
            "Part Number": 1.1, "Part Type": 1.9, "Qty": 0.6,
            "InStk": 1.5, "Donaldson Interchange": 1.3, "In Stock": 1.5,
        }
        for i, c in enumerate(cols):
            w = width_map.get(c)
            if w:
                for cell in [r.cells[i] for r in table.rows]:
                    cell.width = Inches(w)

        out = io.BytesIO(); doc.save(out); out.seek(0)
        return out.read()

    # ---------- Controls ----------
    assets    = sorted(df_ctrl[COL_ASSET].dropna().astype(str).unique().tolist())
    services  = sorted(df_proc[COL_SERVICE].dropna().astype(str).unique().tolist())
    locations = sorted(df_ctrl[COL_LOC_LIST].dropna().astype(str).unique().tolist())

    sel_asset   = st.selectbox("Asset", options=assets, index=0 if assets else None)
    sel_service = st.selectbox("Service", options=services, index=0 if services else None)
    sel_loc     = st.selectbox("Location (Controls[Locations] → Parts_Master[Location2])", options=locations, index=0 if locations else None)

    # ---------- Run ----------
    if st.button("Run Filter"):
        if not assets or not services or not locations:
            st.warning("Workbook appears to be missing required values.")
            st.stop()

        row = df_ctrl.loc[df_ctrl[COL_ASSET].astype(str) == str(sel_asset)]
        if row.empty:
            st.warning(f"No Serial in Controls for asset: {sel_asset}")
            st.stop()
        serial = str(row.iloc[0][COL_SERIAL])

        mask = (
            df_proc[COL_SERIAL].astype(str).str.upper().eq(serial.upper()) &
            df_proc[COL_SERVICE].astype(str).str.upper().eq(sel_service.upper())
        )
        proc_filt = df_proc.loc[mask].copy()
        if proc_filt.empty:
            st.warning(f"No rows for Serial {serial} / Service {sel_service}.")
            st.stop()

        result = unpivot_to_task_then_parts(proc_filt)

        # build PN normalization for matching
        def normalize_pn(pn):
            if pd.isna(pn): return ""
            return "".join(ch for ch in str(pn).upper() if ch.isalnum())
        result["_PN_NORM"] = result["Part Number"].apply(normalize_pn)

        # ---- CAT → Donaldson cross reference ----
        if not df_xref.empty and {XREF_COL_CAT, XREF_COL_DON}.issubset(df_xref.columns):
            xref = df_xref.copy()
            xref["_CAT_NORM"] = xref[XREF_COL_CAT].apply(normalize_pn)
            xref["_DON_RAW"]  = xref[XREF_COL_DON].astype(str)
            result = result.merge(
                xref[["_CAT_NORM", "_DON_RAW"]].drop_duplicates("_CAT_NORM"),
                how="left", left_on="_PN_NORM", right_on="_CAT_NORM"
            )
            def map_don(row_):
                if row_["_PN_NORM"] == "":
                    return ""
                val = row_.get("_DON_RAW")
                if pd.isna(val) or str(val).strip() == "":
                    return "No Interchange"
                return str(val)
            result["Donaldson Interchange"] = result.apply(map_don, axis=1)
            result.drop(columns=["_CAT_NORM", "_DON_RAW"], inplace=True, errors="ignore")
        else:
            result["Donaldson Interchange"] = ""

        # ---- Inventory lookups (Name-first, then Part Numbers; token-based) ----
        def compute_instk(pn_norm: str) -> str:
            if pn_norm == "":
                return ""
            hit = excel_like_first_match(df_inv, pn_norm, sel_loc)
            return inv_text_from_row(hit)

        result["InStk"] = result["_PN_NORM"].apply(compute_instk)

        def compute_instock(di: str) -> str:
            if not di or di.strip() == "" or di.strip().upper() == "NO INTERCHANGE":
                return ""
            di_norm = normalize_pn(di)
            hit = excel_like_first_match(df_inv, di_norm, sel_loc)
            return inv_text_from_row(hit)

        result["In Stock"] = result["Donaldson Interchange"].apply(compute_instock)
        result.drop(columns=["_PN_NORM"], inplace=True, errors="ignore")

        base_cols = [COL_TASKNO, COL_TASK, "Part Number", "Part Type", "Qty",
                    "InStk", "Donaldson Interchange", "In Stock"]
        ordered = [c for c in base_cols if c in result.columns] + [c for c in result.columns if c not in base_cols]

        st.subheader("Filtered Result")
        st.dataframe(result[ordered], use_container_width=True)

        # Downloads
        def to_excel_bytes(df: pd.DataFrame) -> bytes:
            bio = io.BytesIO()
            with pd.ExcelWriter(bio, engine="xlsxwriter") as w:
                df.to_excel(w, index=False, sheet_name="Filtered")
            bio.seek(0)
            return bio.getvalue()

        csv_bytes  = result[ordered].to_csv(index=False).encode("utf-8")
        xlsx_bytes = to_excel_bytes(result[ordered])

        if DOCX_AVAILABLE:
            docx_bytes = to_docx_bytes_sp(
                result[ordered],
                asset=sel_asset, serial=serial, service=sel_service, location=sel_loc
            )

        cols = 4 if (DOCX_AVAILABLE and SHOW_PRINT) else (3 if DOCX_AVAILABLE else 2)
        c = st.columns(cols)
        with c[0]:
            st.download_button("⬇️ CSV", data=csv_bytes,
                            file_name=f"Filtered_{sel_asset}_{sel_service.replace(' ','_')}.csv",
                            mime="text/csv")
        with c[1]:
            st.download_button("⬇️ Excel", data=xlsx_bytes,
                            file_name=f"Filtered_{sel_asset}_{sel_service.replace(' ','_')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        if DOCX_AVAILABLE:
            with c[2]:
                st.download_button("⬇️ Word", data=docx_bytes,
                                file_name=f"Filtered_{sel_asset}_{sel_service.replace(' ','_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            if SHOW_PRINT:
                with c[3]:
                    if st.button("🖨️ Print Word (default printer)"):
                        # Print via temporary file on Windows
                        try:
                            tmp_dir = tempfile.gettempdir()
                            tmp_path = os.path.join(tmp_dir, f"SPF_{int(time.time())}_Filtered.docx")
                            with open(tmp_path, "wb") as f:
                                f.write(docx_bytes)
                            os.startfile(tmp_path, "print")  # type: ignore[attr-defined]
                            st.success(f"Sent to default printer: {os.path.basename(tmp_path)}")
                        except Exception as e:
                            st.error(f"Print failed: {e}")

        st.caption("Task # shows only on header rows. Inventory uses Parts_Master filtered by Location2 = selected Location.")

# ---------- App ----------
st.sidebar.caption(f"SPF Work Orders — v{APP_VERSION}")

cfg = load_config()
cfg = to_plain(cfg)

# Auth
cookie_cfg = cfg.get("cookie", {})
auth = stauth.Authenticate(
    cfg.get("credentials", {}),
    cookie_cfg.get("name", "spf_workorders_portal"),
    cookie_cfg.get("key", "super_secret_key"),
    cookie_cfg.get("expiry_days", 7),
)
name, auth_status, username = auth.login("Login", "main")

if auth_status is False:
    st.error("Username/password is incorrect")
elif auth_status is None:
    st.info("Please log in.")
else:
    auth.logout("Logout", "sidebar")
    st.sidebar.success(f"Logged in as {name}")

    updated = get_data_last_updated_local(cfg)
    if updated: st.sidebar.caption(updated)

    if st.sidebar.button("🔄 Refresh data"):
        st.cache_data.clear()
        st.rerun()

    page = st.sidebar.radio(
        "Page",
        ["🔎 Asset History", "📋 Work Orders", "🔁 Cross Reference", "🧾 Service Report", "📚 Service History", "🧰 Service Procedure Filter"],
        index=1
    )

    # Load workbook bytes + expose Parquet cache tools
    xlsx_path = _xlsx_path_from_cfg(cfg)
    try:
        xlsx_bytes = get_local_xlsx_bytes(cfg)
    except Exception as e:
        st.error(f"Could not load Excel: {e}")
        st.stop()

    parquet_button_and_refresh(xlsx_path, xlsx_bytes)
    xmtime = _mtime(xlsx_path)

    # Access control: Locations from Asset_Master
    try:
        df_am = load_asset_master_df(xmtime, str(xlsx_path))
    except Exception as e:
        st.error(f"Failed to read Asset_Master: {e}")
        st.stop()

    username_ci = str(username).casefold()
    admins_ci = {str(u).casefold() for u in (cfg.get("access", {}).get("admin_usernames", []) or [])}
    is_admin = username_ci in admins_ci
    ul_raw = (cfg.get("access", {}).get("user_locations", {}) or {})
    ul_map_ci = {str(k).casefold(): v for k, v in ul_raw.items()}
    allowed_cfg = ul_map_ci.get(username_ci, [])
    if isinstance(allowed_cfg, str): allowed_cfg = [allowed_cfg]
    allowed_cfg = [a for a in (allowed_cfg or [])]
    star = any(str(a).strip() == "*" for a in allowed_cfg)

    all_locations = sorted(df_am["Location"].dropna().unique().tolist())
    allowed_locations = set(all_locations) if (is_admin or star) else {loc for loc in all_locations if loc in set(allowed_cfg)}
    allowed_norms = {_norm_key(x) for x in allowed_locations}

    # ========= Asset History =========
    if page == "🔎 Asset History":
        st.markdown("### Asset History")
        c1, c2 = st.columns([2, 3])
        with c1:
            loc_options = sorted(allowed_locations)
            chosen_loc = st.selectbox("Location", options=loc_options, index=0, label_visibility="collapsed")
        with c2:
            assets_for_loc = sorted(df_am.loc[df_am["Location"] == chosen_loc, "ASSET"].dropna().unique().tolist())
            chosen_asset = st.selectbox("Asset", options=assets_for_loc, index=0 if assets_for_loc else None, label_visibility="collapsed")

        if not assets_for_loc:
            st.info("No assets for this Location.")
            st.stop()

        try:
            df_all = load_workorders_df(xmtime, str(xlsx_path))
        except Exception as e:
            st.error(f"Failed to read Workorders (history): {e}")
            st.stop()

        df = df_all[(df_all["Location"] == chosen_loc) & (df_all["ASSET"] == chosen_asset)].copy()
        if "QUANTITY RECEIVED" in df.columns and "P/N" in df.columns:
            qnum = pd.to_numeric(df["QUANTITY RECEIVED"], errors="coerce")
            is_part = df["P/N"].astype(str).str.strip().ne("")
            df = df[~(is_part & qnum.notna() & (qnum <= 0))].copy()

        df["__row"] = range(len(df))
        if OPTIONAL_SORT_COL in df.columns:
            df["__sort_key"] = pd.to_numeric(df[OPTIONAL_SORT_COL], errors="coerce").fillna(1).astype(int)
        else:
            df["__sort_key"] = 1
        df.sort_values(by=["WORKORDER","__sort_key","__row"], ascending=[True, True, True], inplace=True)
        df.loc[df["__sort_key"].isin([2, 3]), "WORKORDER"] = ""
        drop_cols = ["__row","__sort_key", OPTIONAL_SORT_COL]
        df_out = df.drop(columns=[c for c in drop_cols if c in df.columns], errors="ignore")

        st.dataframe(df_out, use_container_width=True, hide_index=True)

        c1, c2, _ = st.columns([1, 1, 6])
        with c1:
            st.download_button(
                label="⬇️ Excel (.xlsx)",
                data=to_xlsx_bytes(df_out, sheet="Workorders"),
                file_name=f"WorkOrders_{chosen_loc}_{chosen_asset}.xlsx".replace(" ", "_"),
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        with c2:
            st.download_button(
                label="⬇️ Word (.docx)",
                data=to_docx_bytes(df_out, title=f"Work Orders — {chosen_loc} — {chosen_asset}"),
                file_name=f"WorkOrders_{chosen_loc}_{chosen_asset}.docx".replace(" ", "_"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        st.stop()

    # ========= Work Orders =========
    if page == "📋 Work Orders":
        st.markdown("### Work Orders — Filtered Views (flags from workbook)")

        try:
            df_master = load_wo_master_df(xmtime, str(xlsx_path))
        except Exception as e:
            st.error(f"Failed to read '{SHEET_WO_MASTER}': {e}")
            st.stop()

        opt_users = load_users_sheet(xmtime, str(xlsx_path))
        df_master = df_master[df_master["Location"].isin(allowed_locations)].copy()
        total_in_scope = len(df_master)

        c1, c2, c3, c4 = st.columns([2, 2, 2, 3])
        with c1:
            loc_values = sorted(df_master["Location"].dropna().unique().tolist())
            loc_all_label = f"« All my locations ({len(loc_values)}) »"
            chosen_loc = st.selectbox("Location", options=[loc_all_label] + loc_values, index=0, label_visibility="collapsed")

        df_scope = df_master if chosen_loc == loc_all_label else df_master[df_master["Location"] == chosen_loc].copy()

        with c2:
            if opt_users is not None:
                user_choices = ["— Any user —"] + opt_users
            else:
                derived_users = sorted([u for u in df_scope.get("Assigned to", pd.Series([], dtype=str)).dropna().astype(str).str.strip().unique().tolist() if u])
                user_choices = ["— Any user —"] + derived_users
            sel_user = st.selectbox("Assigned user", options=user_choices, index=0, label_visibility="collapsed")

        with c3:
            raw_teams = df_scope.get("Teams Assigned to", pd.Series([], dtype=str)).fillna("").astype(str)
            token_set = set()
            for v in raw_teams:
                for t in re.split(r"[;,]", v):
                    t = t.strip()
                    if t:
                        token_set.add(t)
            team_opts = ["— Any team —"] + sorted(token_set)
            sel_team = st.selectbox("Team", options=team_opts, index=0, label_visibility="collapsed")

        with c4:
            view = st.radio("View", ["All","Open","Overdue","Scheduled (Planning)","Completed","Old"],
                            horizontal=True, index=1 if "IsOpen" in df_scope.columns else 0)

        if sel_user != "— Any user —":
            df_scope = df_scope[df_scope["Assigned to"].astype(str).str.strip() == sel_user].copy()
            if df_scope.empty:
                st.warning("User not found at this location (or no work orders match).")
                st.dataframe(df_scope, use_container_width=True, hide_index=True)
                st.stop()

        if sel_team != "— Any team —":
            def team_hit(s: str) -> bool:
                if not s: return False
                parts = {p.strip() for p in re.split(r"[;,]", str(s)) if p.strip()}
                return sel_team.strip() in parts
            df_scope = df_scope[df_scope["Teams Assigned to"].fillna("").astype(str).map(team_hit)].copy()

        def pick_view(df_in: pd.DataFrame) -> pd.DataFrame:
            if view == "All" or not {"IsOpen","IsOverdue","IsScheduled","IsCompleted","IsOld"}.issubset(df_in.columns):
                return df_in
            col = {"Open":"IsOpen","Overdue":"IsOverdue","Scheduled (Planning)":"IsScheduled","Completed":"IsCompleted","Old":"IsOld"}[view]
            return df_in[df_in[col]].copy()

        df_view = pick_view(df_scope)

        def present(cols: list[str]) -> list[str]:
            return [c for c in cols if c in df_view.columns]

        cols_all       = present(["ID","Title","Description","Asset","Created on","Planned Start Date","Due date","Started on","Completed on","Assigned to","Teams Assigned to","Location"])
        cols_open      = present(["ID","Title","Description","Asset","Created on","Due date","Assigned to","Teams Assigned to","Location"])
        cols_overdue   = present(["ID","Title","Description","Asset","Due date","Assigned to","Teams Assigned to","Location"])
        cols_sched     = present(["ID","Title","Description","Asset","Planned Start Date","Due date","Assigned to","Teams Assigned to","Location"])
        cols_completed = present(["ID","Title","Description","Asset","Completed on","Assigned to","Teams Assigned to","Location"])
        cols_old       = present(["ID","Title","Description","Asset","Created on","Due date","Completed on","Assigned to","Teams Assigned to","Location"])

        if view == "Open":
            use_cols = cols_open or df_view.columns.tolist()
            sort_keys = [k for k in ["Due date","Created on","Title"] if k in df_view.columns]
        elif view == "Overdue":
            use_cols = cols_overdue or df_view.columns.tolist()
            sort_keys = [k for k in ["Due date","Title"] if k in df_view.columns]
        elif view == "Scheduled (Planning)":
            use_cols = cols_sched or df_view.columns.tolist()
            sort_keys = [k for k in ["Planned Start Date","Due date","Title"] if k in df_view.columns]
        elif view == "Completed":
            use_cols = cols_completed or df_view.columns.tolist()
            sort_keys = [k for k in ["Completed on","Title"] if k in df_view.columns]
        elif view == "Old":
            use_cols = cols_old or df_view.columns.tolist()
            sort_keys = [k for k in ["Created on","Due date","Title"] if k in df_view.columns]
        else:
            use_cols = cols_all or df_view.columns.tolist()
            sort_keys = [k for k in ["Completed on","Due date","Planned Start Date","Created on","Title"] if k in df_view.columns]

        if sort_keys: df_view = df_view.sort_values(by=sort_keys, na_position="last")

        st.caption(f"In scope: {total_in_scope}  •  After location/user/team filters: {len(df_scope)}  •  Showing ({view}): {len(df_view)}")
        st.dataframe(df_view[use_cols], use_container_width=True, hide_index=True)

        with st.expander("🗓️ 7-day Scheduled Planner (printable)", expanded=False):
            df_sched = df_scope[df_scope.get("IsScheduled", False)].copy()
            if df_sched.empty:
                st.info("No scheduled items.")
            else:
                df_sched["Planned Start Date"] = pd.to_datetime(df_sched["Planned Start Date"], errors="coerce")
                start_base = pd.to_datetime(datetime.now().date())
                dates = [start_base + timedelta(days=i) for i in range(7)]
                labels = [d.strftime("%a %m/%d") for d in dates]
                cols = st.columns(7)
                for i, d in enumerate(dates):
                    with cols[i]:
                        st.markdown(f"**{labels[i]}**")
                        day_rows = df_sched[df_sched["Planned Start Date"].dt.date == d.date()]
                        if day_rows.empty:
                            st.caption("—")
                        else:
                            show_cols = [c for c in ["ID","Title","Description","Asset","Assigned to","Location"] if c in day_rows.columns]
                            st.dataframe(day_rows[show_cols], use_container_width=True, hide_index=True)

        c1, c2, _ = st.columns([1, 1, 6])
        with c1:
            st.download_button("⬇️ Excel (.xlsx)", data=to_xlsx_bytes(df_view[use_cols], sheet="WorkOrders"),
                               file_name=f"WorkOrders_{view.replace(' ','_')}.xlsx",
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        with c2:
            st.download_button("⬇️ Word (.docx)", data=to_docx_bytes(df_view[use_cols], title=f"Work Orders — {view}"),
                               file_name=f"WorkOrders_{view.replace(' ','_')}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.stop()

    # ========= Cross Reference =========
    if page == "🔁 Cross Reference":
        st.markdown("### Cross-Reference Finder")
        xref_debug = st.sidebar.checkbox("XRef: show detected headers", value=False)

        # Load local workbook (no GitHub/secrets)
        try:
            inv_df = xref_load_inventory_df(cfg, debug=xref_debug)
            merge_df, long_df = xref_load_merge_wide_and_long(cfg, debug=xref_debug)
        except Exception as e:
            st.error(f"Cross-Reference: failed to read local file: {e}")
            st.stop()

        # Optional Location2 filter (from inventory)
        loc_options = []
        if XREF_INV_COL_LOC2 in inv_df.columns:
            loc_options = sorted(inv_df[XREF_INV_COL_LOC2].dropna().astype(str).unique().tolist())
        use_loc = st.sidebar.checkbox("Filter by a specific Location2", value=True)
        sel_loc2 = st.sidebar.selectbox("Location2", options=loc_options, index=0) if use_loc and loc_options else None

        pn_input = st.text_input("Enter a part number (any brand)", value="", placeholder="e.g., 1R-0716 or P554005")
        if st.button("Find"):
            if not pn_input.strip():
                st.warning("Please enter a part number."); st.stop()

            # 1) Resolve rowset + primary brand
            row_ids, primary_brand = xref_resolve_rowset_and_primary_brand(pn_input, long_df)
            if not row_ids:
                st.error("No cross-reference row found for that part number.")
                st.stop()

            # 2) Build xrefs from those rows
            xrefs = xref_build_crossrefs_from_rows(merge_df, row_ids)

            # 3) Show detected manufacturer and place that brand first (if any)
            if primary_brand:
                st.markdown(f"**Manufacturer detected:** {primary_brand}")
                xrefs = pd.concat(
                    [xrefs[xrefs["Brand"] == primary_brand], xrefs[xrefs["Brand"] != primary_brand]],
                    ignore_index=True
                )
            else:
                st.markdown("**Manufacturer detected:** _Unknown_")

            # 4) Ensure the exact typed PN is included for stock search
            xrefs = xref_add_entered_pn_to_xrefs(xrefs, pn_input, primary_brand)

            # 5) Show cross-refs
            st.subheader("Cross-References")
            st.dataframe(xrefs[["Brand","PartNumber"]], use_container_width=True)

            # 6) Per-xref inventory search; aggregate per inventory row (_ROWID)
            st.subheader("Inventory matches")
            hits = xref_inventory_lookup_per_xref(inv_df, xrefs, sel_loc2 if use_loc else None)
            if hits.empty:
                st.info("No stock at the selected Location2 for any cross references." if use_loc else "No stock found for any cross references.")
            else:
                st.dataframe(hits, use_container_width=True)
        st.stop()

    # ========= Service Report =========
    if page == "🧾 Service Report":
        st.markdown("### Service Report")

        raw_sr, canon_sr, source_sheet = load_service_report_df(xmtime, str(xlsx_path))
        if raw_sr is None:
            st.warning("No 'Service Report' sheet found.")
            st.stop()

        # Location filter (shorter)
        loc_col = None
        for c in raw_sr.columns:
            if c.strip().lower() in {"location","ns location","location2"}:
                loc_col = c; break

        if loc_col:
            loc_values_all = raw_sr[loc_col].astype(str)
            loc_candidates = sorted({v for v in loc_values_all if _norm_key(v) in allowed_norms})
        else:
            loc_candidates = sorted(allowed_locations)

        c_loc, _ = st.columns([3, 7])
        with c_loc:
            loc_all_label = f"« All my locations ({len(loc_candidates)}) »" if loc_candidates else "« All my locations »"
            chosen_loc = st.selectbox("Location", options=[loc_all_label] + loc_candidates if loc_candidates else [loc_all_label],
                                      index=0, label_visibility="collapsed")

        if loc_col and chosen_loc != loc_all_label:
            mask_norm = loc_values_all.map(_norm_key) == _norm_key(chosen_loc)
            raw_show = raw_sr[mask_norm].copy()
            if "Location" in canon_sr.columns:
                canon_in_scope = canon_sr[canon_sr["Location"].map(_norm_key) == _norm_key(chosen_loc)].copy()
            else:
                canon_in_scope = canon_sr.copy()
        else:
            raw_show = raw_sr.copy()
            canon_in_scope = canon_sr.copy()

        t_report, t_due, t_over = st.tabs(["Report", "Coming Due", "Overdue"])

        # --- Report (as-is, minus Schedule/Today) with Date normalized to date-only ---
        with t_report:
            drop_cols = [c for c in ["Schedule","Today"] if c in raw_show.columns]
            show_rep = raw_show.drop(columns=drop_cols) if drop_cols else raw_show
            # Normalize a column explicitly named 'Date' to YYYY-MM-DD for display only
            if "Date" in show_rep.columns:
                show_rep = show_rep.copy()
                show_rep["Date"] = show_rep["Date"].map(_norm_date_any)
            st.caption(f"Source: {source_sheet}  •  Rows: {len(show_rep)}  •  No filters other than Location.")
            st.dataframe(show_rep, use_container_width=True, hide_index=True)
            c1, c2, _ = st.columns([1,1,6])
            with c1:
                st.download_button("⬇️ Excel (.xlsx)", data=to_xlsx_bytes(show_rep, sheet="Service_Report"),
                                   file_name="Service_Report.xlsx",
                                   mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            with c2:
                st.download_button("⬇️ Word (.docx)", data=to_docx_bytes(show_rep, title="Service Report — As Is"),
                                   file_name="Service_Report.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Helper: compute Coming Due / Overdue frames (vectorized thresholds)
        def _due_frames(df_can: pd.DataFrame):
            if df_can is None or df_can.empty:
                return pd.DataFrame(), pd.DataFrame()
            df = df_can.copy()
            # 5% threshold if meter type mentions miles, else 10%
            mt = df.get("__MeterType_norm")
            thr_series = pd.Series(0.10, index=df.index)
            if mt is not None:
                thr_series = thr_series.where(~mt.astype(str).str.contains("mile"), 0.05)

            conds = []
            condA = (
                df["__Schedule_num"].notna() &
                (pd.to_numeric(df["__Schedule_num"], errors="coerce") > 0) &
                df["__Remaining_num"].notna() &
                (pd.to_numeric(df["__Remaining_num"], errors="coerce") >= 0)
            )
            if condA.any():
                condA2 = pd.to_numeric(df["__Remaining_num"], errors="coerce") <= (pd.to_numeric(df["__Schedule_num"], errors="coerce") * thr_series)
                conds.append(condA & condA2)
            if "__PctRemain_num" in df.columns:
                condB = df["__PctRemain_num"].notna()
                if condB.any():
                    condB2 = pd.to_numeric(df["__PctRemain_num"], errors="coerce") <= thr_series
                    conds.append(condB & condB2)
            coming_due_mask = pd.Series(False, index=df.index)
            for c in conds: coming_due_mask |= c
            coming_due = df[coming_due_mask].copy()

            today = pd.Timestamp.today().normalize()
            overdue_mask = pd.Series(False, index=df.index)
            if "__Remaining_num" in df.columns:
                overdue_mask |= (pd.to_numeric(df["__Remaining_num"], errors="coerce") < 0)
            if "__Due_dt" in df.columns:
                overdue_mask |= (pd.to_datetime(df["__Due_dt"], errors="coerce") < today)
            overdue = df[overdue_mask].copy()
            return coming_due, overdue

        coming_due_df, overdue_df = _due_frames(canon_in_scope)

        def reorder_with_optional(df: pd.DataFrame, pref: list[str]) -> pd.DataFrame:
            cols = [c for c in pref if c in df.columns]
            rest = [c for c in df.columns if c not in cols]
            return df[cols + rest]

        with t_due:
            st.caption("Coming Due = Remaining ≤ 10% of Schedule (or ≤ 5% if Meter Type contains 'miles').")
            if coming_due_df.empty:
                st.info("No items are coming due based on the available columns.")
            else:
                base = ["WO_ID","Title","Service","Asset","Location","Date","User","Status",
                        "Schedule","Remaining","Percent Remaining","Meter Type","Due Date","Notes","MReading"]
                show = coming_due_df[[c for c in base if c in coming_due_df.columns]].copy()
                show = show.sort_values(by=[c for c in ["Due Date","Percent Remaining","Remaining"] if c in show.columns], na_position="last")
                st.dataframe(show, use_container_width=True, hide_index=True)
                c1, c2, _ = st.columns([1,1,6])
                with c1:
                    st.download_button("⬇️ Excel (.xlsx)", data=to_xlsx_bytes(show, sheet="Service_Coming_Due"),
                                       file_name="Service_Coming_Due.xlsx",
                                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                with c2:
                    st.download_button("⬇️ Word (.docx)", data=to_docx_bytes(show, title="Service — Coming Due"),
                                       file_name="Service_Coming_Due.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        with t_over:
            st.caption("Overdue = Remaining < 0, or Due Date earlier than today.")
            if overdue_df.empty:
                st.info("No overdue items found.")
            else:
                base = ["WO_ID","Title","Service","Asset","Location","Date","User","Status",
                        "Schedule","MReading","Remaining","Percent Remaining","Meter Type","Due Date","Notes"]
                show = overdue_df[[c for c in base if c in overdue_df.columns]].copy()
                show = reorder_with_optional(show, base)
                show = show.sort_values(by=[c for c in ["Due Date","Remaining"] if c in show.columns], na_position="last")
                st.dataframe(show, use_container_width=True, hide_index=True)
                c1, c2, _ = st.columns([1,1,6])
                with c1:
                    st.download_button("⬇️ Excel (.xlsx)", data=to_xlsx_bytes(show, sheet="Service_Overdue"),
                                       file_name="Service_Overdue.xlsx",
                                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                with c2:
                    st.download_button("⬇️ Word (.docx)", data=to_docx_bytes(show, title="Service — Overdue"),
                                       file_name="Service_Overdue.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.stop()

    # ========= Service History =========
    if page == "📚 Service History":
        st.markdown("### Service History")

        df_hist, used_sheet_or_err = load_service_history_df(xmtime, str(xlsx_path))
        if df_hist is None or df_hist.empty:
            msg = used_sheet_or_err or "unknown error"
            st.warning(f"No Service History data found. Tried: {SHEET_WO_SERVICE_CANDS}. Last error: {msg}")
            st.stop()

        # Normalize and restrict to allowed by Location2
        if "Location2" in df_hist.columns:
            df_hist["__LocNorm"] = df_hist["Location2"].map(_norm_key)
            df_hist = df_hist[df_hist["__LocNorm"].isin(allowed_norms)].copy()

        # Filters: Location + Asset (compact)
        c1, c2 = st.columns([2, 3])
        with c1:
            if "Location2" in df_hist.columns:
                loc_values = sorted(df_hist["Location2"].dropna().unique().tolist())
            else:
                loc_values = []
            loc_all_label = f"« All my locations ({len(loc_values)}) »" if loc_values else "« All my locations »"
            chosen_loc = st.selectbox("Location", options=[loc_all_label] + loc_values if loc_values else [loc_all_label],
                                      index=0, label_visibility="collapsed")

        if chosen_loc != loc_all_label and "Location2" in df_hist.columns:
            scope = df_hist[df_hist["Location2"].map(_norm_key) == _norm_key(chosen_loc)].copy()
        else:
            scope = df_hist.copy()

        with c2:
            assets = sorted([a for a in scope.get("Asset", pd.Series([], dtype=str)).dropna().astype(str).str.strip().unique().tolist() if a])
            sel_asset = st.selectbox("Asset", options=assets, index=0 if assets else None, label_visibility="collapsed")

        if not assets:
            st.info("No assets available in this Location.")
            st.stop()

        scope = scope[scope["Asset"] == sel_asset] if "Asset" in scope.columns else scope

        if "Date" in scope.columns:
            scope = scope.copy()
            scope["__Date_dt"] = pd.to_datetime(scope["Date"], errors="coerce")
            scope = scope.sort_values(by="__Date_dt", ascending=False, na_position="last").drop(columns="__Date_dt")

        # Show MReading right after Service
        col_order = [c for c in ["Date","WO_ID","Title","Service","MReading","MHours","Asset","User","Location2","Notes","Status"] if c in scope.columns]
        st.caption(f"Sheet used: {used_sheet_or_err} • Rows: {len(scope)}")
        st.dataframe(scope[col_order] if col_order else scope, use_container_width=True, hide_index=True)

        c1, c2, _ = st.columns([1,1,6])
        with c1:
            st.download_button(
                "⬇️ Excel (.xlsx)",
                data=to_xlsx_bytes(scope[col_order] if col_order else scope, sheet="Service_History"),
                file_name=f"Service_History_{sel_asset.replace(' ','_')}.xlsx" if assets else "Service_History.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        with c2:
            st.download_button(
                "⬇️ Word (.docx)",
                data=to_docx_bytes(scope[col_order] if col_order else scope, title=f"Service History — {sel_asset}" if assets else "Service History"),
                file_name=f"Service_History_{sel_asset.replace(' ','_')}.docx" if assets else "Service_History.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        st.stop()

    # ========= Service Procedure Filter =========
    if page == "🧰 Service Procedure Filter":
        render_service_procedure_page(xmtime, str(xlsx_path))
        st.stop()
